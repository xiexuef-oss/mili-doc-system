"""
GJB/Z 299D-2024 表格数据提取器。
从已修正的 DOCX 中提取 834 张表格的数值数据，
输出 JSON 文件供 Gjb299dDataImporter 批量导入。

用法: python extract_299d_tables.py
输出: 299d_cache_import.json
"""

from docx import Document
import json
import re
import os

DOCX_PATH = r"D:\成都天奥信息科技有限公司\GJB\可靠性相关\（已压缩）GJBZ 299D-2024\（已压缩）GJBZ 299D-2024电子设备可靠性预计手册_已修正.docx"
OUTPUT_PATH = os.path.join(os.path.dirname(__file__), "299d_cache_import.json")

def extract_all_tables(docx_path):
    """Extract all tables from DOCX with section context."""
    doc = Document(docx_path)
    
    # First pass: build section → paragraph index mapping
    current_section = "前言"
    current_subsection = ""
    para_sections = {}
    
    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        if not text:
            continue
        
        # Detect section headings
        # Patterns: "5.1 微电路", "5.1.1 概述", "5.3.2 普通二极管"
        sec_match = re.match(r'^(?:([\d]+(?:\.[\d]+)*)\s+)(.+)', text)
        if sec_match:
            num = sec_match.group(1)
            title = sec_match.group(2)
            depth = num.count('.')
            if depth == 0:  # Top level: "5 工作状态..."
                current_section = f"{num} {title}"
                current_subsection = ""
            elif depth == 1:  # "5.1 微电路"
                current_subsection = f"{num} {title}"
            elif depth == 2:  # "5.1.1 概述"
                current_subsection = f"{num} {title}"
        
        para_sections[i] = (current_section, current_subsection)
    
    # Second pass: extract tables and associate with nearest section
    entries = []
    
    for ti, table in enumerate(doc.tables):
        # Find nearest preceding paragraph with section info
        sec = "未知"
        subsec = ""
        
        # Walk backwards to find section
        # Tables in DOCX are interleaved with paragraphs in document order
        # Approximate by using table index as position
        for pi in sorted(para_sections.keys(), reverse=True):
            if pi <= ti * 2:  # rough heuristic
                sec, subsec = para_sections[pi]
                break
        
        # Extract table data
        rows = len(table.rows)
        cols = len(table.columns)
        
        table_info = {
            "table_index": ti,
            "rows": rows,
            "cols": cols,
            "section": sec,
            "subsection": subsec,
            "headers": [],
            "data_rows": []
        }
        
        # Extract header (first 1-3 rows)
        for ri in range(min(3, rows)):
            row_data = []
            for ci in range(cols):
                cell_text = table.rows[ri].cells[ci].text.strip()
                row_data.append(cell_text)
            table_info["headers"].append(row_data)
        
        # Extract data (first 5 data rows for preview)
        for ri in range(3, min(8, rows)):
            row_data = []
            for ci in range(cols):
                cell_text = table.rows[ri].cells[ci].text.strip()
                row_data.append(cell_text)
            table_info["data_rows"].append(row_data)
        
        entries.append(table_info)
        
        if ti < 50 or ti % 50 == 0:
            print(f"  Table {ti}: {rows}×{cols} | {subsec[:40] if subsec else sec[:40]}")
    
    return entries

def categorize_tables(entries):
    """Categorize tables based on section content and structure."""
    categories = {
        "微电路": [],
        "声表面波器件": [],
        "半导体分立器件": [],
        "光电子器件": [],
        "真空电子器件": [],
        "电阻器": [],
        "敏感电阻器": [],
        "电位器": [],
        "电容器": [],
        "感性元件": [],
        "继电器": [],
        "开关": [],
        "电连接器": [],
        "微特电机": [],
        "印制板": [],
        "磁性器件": [],
        "振荡器": [],
        "滤波器": [],
        "电池": [],
        "灯": [],
        "激光器": [],
        "压电陀螺": [],
        "光纤连接器": [],
        "其他": [],
        "附录": [],
        "引用文件": [],
        "概述/术语": [],
    }
    
    for entry in entries:
        subsec = entry["subsection"] + entry["section"]
        matched = False
        for cat in categories:
            if cat in subsec:
                categories[cat].append(entry)
                matched = True
                break
        if not matched:
            categories["其他"].append(entry)
    
    return categories


if __name__ == "__main__":
    print(f"Reading: {DOCX_PATH}")
    print("Extracting tables...")
    
    entries = extract_all_tables(DOCX_PATH)
    
    print(f"\nTotal tables: {len(entries)}")
    
    # Categorize
    categories = categorize_tables(entries)
    
    print("\n=== 表格分类统计 ===")
    for cat, tabs in sorted(categories.items(), key=lambda x: -len(x[1])):
        if tabs:
            print(f"  {cat}: {len(tabs)} 张表")
    
    # Sample key tables for manual verification
    print("\n=== 关键表格采样 ===")
    for cat in ["微电路", "半导体分立器件", "电阻器", "电容器"]:
        tabs = categories.get(cat, [])
        if tabs:
            print(f"\n--- {cat} (前 3 张表) ---")
            for t in tabs[:3]:
                print(f"  表{t['table_index']}: {t['rows']}×{t['cols']}")
                if t['headers']:
                    print(f"    表头: {t['headers'][0][:5]}")
                if t['data_rows']:
                    print(f"    数据行1: {t['data_rows'][0][:5]}")
    
    # Save to JSON
    with open(OUTPUT_PATH, 'w', encoding='utf-8') as f:
        json.dump(entries, f, ensure_ascii=False, indent=2)
    
    print(f"\nSaved {len(entries)} tables to: {OUTPUT_PATH}")
